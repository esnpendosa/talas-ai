"""
TALAS AI — Report Generator (Phase 14)
Generate laporan telaah regulasi dalam format DOCX, PDF, atau JSON.

Struktur laporan:
  I. IDENTITAS
  II. LATAR BELAKANG
  III. DASAR HUKUM
  IV. MATERI MUATAN
  V. HASIL ANALISIS
  VI. POTENSI PERMASALAHAN
  VII. REKOMENDASI
  VIII. KESIMPULAN

PRINSIP:
- Disclaimer wajib di setiap halaman: "TINJAUAN AWAL AI — WAJIB VERIFIKASI MANUSIA."
- Graceful fallback ke JSON jika python-docx/reportlab tidak tersedia
"""
from __future__ import annotations

import json
import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.analysis import Analysis, AnalysisFinding
from app.models.regulation import Regulation
from app.models.report import Report, ReportVersion

logger = logging.getLogger("talas_ai.reports")

DISCLAIMER = "TINJAUAN AWAL AI — WAJIB VERIFIKASI MANUSIA."


async def generate_report(
    db: AsyncSession,
    analysis_id: int,
    format: str = "json",
    generated_by: Optional[int] = None,
) -> Dict:
    """
    Generate laporan telaah regulasi.

    Args:
        db: Database session
        analysis_id: ID analisis
        format: docx | pdf | json (fallback ke json jika lib tidak tersedia)
        generated_by: user_id yang men-generate

    Returns: dict dengan info laporan
    """
    # Ambil data analisis
    result = await db.execute(
        select(Analysis).where(Analysis.id == analysis_id)
    )
    analysis = result.scalar_one_or_none()
    if not analysis:
        raise ValueError(f"Analisis ID {analysis_id} tidak ditemukan.")

    # Ambil data regulasi
    result = await db.execute(
        select(Regulation).where(Regulation.id == analysis.regulation_id)
    )
    regulation = result.scalar_one_or_none()

    # Ambil findings
    result = await db.execute(
        select(AnalysisFinding)
        .where(AnalysisFinding.analysis_id == analysis_id)
        .order_by(AnalysisFinding.finding_type, AnalysisFinding.pasal)
    )
    findings = result.scalars().all()

    # Siapkan konten laporan
    report_data = _build_report_data(analysis, regulation, findings)

    # Buat direktori export
    export_dir = Path(settings.EXPORT_DIR)
    export_dir.mkdir(parents=True, exist_ok=True)

    # Timestamp untuk nama file
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    reg_name = _safe_filename(regulation.judul if regulation else "regulasi")[:30]
    base_filename = f"telaah_{reg_name}_{ts}"

    # Generate file
    actual_format = format.lower()
    file_path = None
    file_size = 0

    if actual_format == "docx":
        try:
            file_path = await _generate_docx(report_data, export_dir, base_filename)
            file_size = os.path.getsize(file_path) if file_path and os.path.exists(file_path) else 0
        except ImportError:
            logger.warning("python-docx tidak tersedia. Fallback ke JSON.")
            actual_format = "json"
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            actual_format = "json"

    if actual_format == "pdf":
        try:
            file_path = await _generate_pdf(report_data, export_dir, base_filename)
            file_size = os.path.getsize(file_path) if file_path and os.path.exists(file_path) else 0
        except ImportError:
            logger.warning("reportlab tidak tersedia. Fallback ke JSON.")
            actual_format = "json"
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            actual_format = "json"

    if actual_format == "json" or file_path is None:
        file_path = export_dir / f"{base_filename}.json"
        content = json.dumps(report_data, ensure_ascii=False, indent=2, default=str)
        file_path.write_text(content, encoding="utf-8")
        file_size = file_path.stat().st_size
        actual_format = "json"

    # Buat record Report
    title = f"Telaah Regulasi: {regulation.judul if regulation else 'Regulasi'}"
    report = Report(
        analysis_id=analysis_id,
        regulation_id=analysis.regulation_id,
        title=title[:500],
        report_type="TELAAH",
        status="FINAL",
        generated_by=generated_by,
    )
    db.add(report)
    await db.flush()

    # Buat ReportVersion
    version = ReportVersion(
        report_id=report.id,
        version_number=1,
        file_format=actual_format,
        file_path=str(file_path),
        file_size=file_size,
        is_current=True,
        generated_by=generated_by,
    )
    db.add(version)
    await db.commit()
    await db.refresh(report)

    return {
        "success": True,
        "report_id": report.id,
        "title": title,
        "format": actual_format,
        "file_path": str(file_path),
        "file_size": file_size,
        "disclaimer": DISCLAIMER,
    }


def _build_report_data(
    analysis: Analysis,
    regulation: Optional[Regulation],
    findings: List[AnalysisFinding],
) -> Dict:
    """Build struktur data laporan."""
    # Pisahkan findings berdasarkan tipe
    legal_basis_findings = [f for f in findings if f.finding_type == "LEGAL_BASIS"]
    conflict_findings = [f for f in findings if f.finding_type == "CONFLICT"]
    consistency_findings = [f for f in findings if f.finding_type == "CONSISTENCY"]

    # Potensi permasalahan = findings dengan status bermasalah
    problem_statuses = {"NOT_FOUND", "POTENTIAL_CONFLICT", "DIFFERENCE", "NEEDS_REVIEW"}
    problem_findings = [f for f in findings if f.status in problem_statuses]

    # Build rekomendasi
    recommendations = []
    for f in findings:
        if f.recommendation:
            recommendations.append(f.recommendation)

    now_str = datetime.now(timezone.utc).isoformat()

    return {
        "disclaimer": DISCLAIMER,
        "generated_at": now_str,
        "sections": {
            "I_IDENTITAS": {
                "judul": regulation.judul if regulation else "Regulasi tidak teridentifikasi",
                "jenis": regulation.jenis if regulation else "-",
                "nomor": regulation.nomor if regulation else "-",
                "tahun": regulation.tahun if regulation else "-",
                "status": regulation.status if regulation else "-",
            },
            "II_LATAR_BELAKANG": {
                "analysis_type": analysis.analysis_type,
                "analysis_id": analysis.id,
                "started_at": str(analysis.started_at) if analysis.started_at else "-",
                "completed_at": str(analysis.completed_at) if analysis.completed_at else "-",
                "ai_provider": analysis.ai_provider or "Mock/Local",
                "catatan": (
                    "Telaah ini dilakukan secara otomatis oleh AI TALAS. "
                    "Hasil merupakan tinjauan awal dan wajib diverifikasi oleh analis hukum."
                ),
            },
            "III_DASAR_HUKUM": {
                "total_pasal_dianalisis": len(legal_basis_findings),
                "dasar_hukum_ditemukan": sum(1 for f in legal_basis_findings if f.status == "FOUND"),
                "dasar_hukum_tidak_ditemukan": sum(1 for f in legal_basis_findings if f.status == "NOT_FOUND"),
                "perlu_review": sum(1 for f in legal_basis_findings if f.status == "NEEDS_REVIEW"),
                "findings": [
                    {
                        "pasal": f.pasal,
                        "status": f.status,
                        "confidence": f.confidence,
                        "finding": f.finding,
                    }
                    for f in legal_basis_findings
                ],
            },
            "IV_MATERI_MUATAN": {
                "total_pasal": analysis.total_articles or len(findings),
                "catatan": "Materi muatan dianalisis berdasarkan dokumen yang diupload.",
            },
            "V_HASIL_ANALISIS": {
                "total_findings": len(findings),
                "by_type": {
                    "LEGAL_BASIS": len(legal_basis_findings),
                    "CONFLICT": len(conflict_findings),
                    "CONSISTENCY": len(consistency_findings),
                },
                "by_status": {
                    status: sum(1 for f in findings if f.status == status)
                    for status in ["FOUND", "NOT_FOUND", "NEEDS_REVIEW", "NO_ISSUE",
                                   "DIFFERENCE", "POTENTIAL_CONFLICT"]
                },
            },
            "VI_POTENSI_PERMASALAHAN": {
                "total": len(problem_findings),
                "items": [
                    {
                        "pasal": f.pasal,
                        "tipe": f.finding_type,
                        "status": f.status,
                        "finding": f.finding,
                    }
                    for f in problem_findings
                ],
            },
            "VII_REKOMENDASI": {
                "items": list(set(recommendations[:20])),  # Deduplicate, max 20
                "catatan": "Rekomendasi di atas adalah saran awal AI dan wajib diverifikasi.",
            },
            "VIII_KESIMPULAN": {
                "disclaimer": DISCLAIMER,
                "kesimpulan": (
                    f"Telaah AI terhadap {analysis.analysis_type} telah selesai. "
                    f"Ditemukan {len(findings)} temuan. "
                    f"{len(problem_findings)} memerlukan perhatian khusus. "
                    "WAJIB diverifikasi oleh analis hukum yang berwenang sebelum digunakan."
                ),
            },
        },
    }


async def _generate_docx(report_data: Dict, export_dir: Path, base_filename: str) -> str:
    """Generate laporan dalam format DOCX menggunakan python-docx."""
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    doc = Document()

    # Title
    title_para = doc.add_heading("LAPORAN TELAAH REGULASI", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Disclaimer header
    disclaimer_para = doc.add_paragraph(DISCLAIMER)
    disclaimer_para.runs[0].bold = True

    sections = report_data.get("sections", {})

    for section_key, section_data in sections.items():
        section_num = section_key.split("_")[0]
        section_name = " ".join(section_key.split("_")[1:])
        doc.add_heading(f"{section_num}. {section_name}", level=1)

        if isinstance(section_data, dict):
            for key, value in section_data.items():
                if key == "findings" or key == "items":
                    continue
                doc.add_paragraph(f"{key}: {value}")

            # Handle nested findings/items
            for list_key in ("findings", "items"):
                items = section_data.get(list_key, [])
                if items and isinstance(items, list):
                    for item in items[:20]:
                        if isinstance(item, dict):
                            text = " | ".join(f"{k}: {v}" for k, v in item.items() if v)
                            p = doc.add_paragraph(text, style="List Bullet")

    # Footer disclaimer di setiap section
    doc.add_paragraph()
    footer = doc.add_paragraph(DISCLAIMER)
    footer.runs[0].bold = True

    file_path = export_dir / f"{base_filename}.docx"
    doc.save(str(file_path))
    return str(file_path)


async def _generate_pdf(report_data: Dict, export_dir: Path, base_filename: str) -> str:
    """Generate laporan dalam format PDF menggunakan reportlab."""
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.lib.units import cm  # type: ignore
    from reportlab.platypus import (  # type: ignore
        SimpleDocTemplate, Paragraph, Spacer, PageBreak
    )

    file_path = export_dir / f"{base_filename}.pdf"

    doc_pdf = SimpleDocTemplate(
        str(file_path),
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph("LAPORAN TELAAH REGULASI", styles["Title"]))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(f"<b>{DISCLAIMER}</b>", styles["Normal"]))
    story.append(Spacer(1, 0.5 * cm))

    sections = report_data.get("sections", {})
    for section_key, section_data in sections.items():
        section_num = section_key.split("_")[0]
        section_name = " ".join(section_key.split("_")[1:])
        story.append(Paragraph(f"{section_num}. {section_name}", styles["Heading1"]))

        if isinstance(section_data, dict):
            for key, value in section_data.items():
                if key in ("findings", "items"):
                    items = section_data[key]
                    if isinstance(items, list):
                        for item in items[:20]:
                            text = str(item)[:300] if not isinstance(item, dict) else \
                                   " | ".join(f"{k}: {v}" for k, v in item.items() if v)
                            story.append(Paragraph(f"• {text}", styles["Normal"]))
                    continue
                story.append(Paragraph(f"<b>{key}:</b> {value}", styles["Normal"]))
        story.append(Spacer(1, 0.3 * cm))

    # Footer disclaimer
    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(f"<b>{DISCLAIMER}</b>", styles["Normal"]))

    doc_pdf.build(story)
    return str(file_path)


def _safe_filename(name: str) -> str:
    """Bersihkan nama untuk digunakan sebagai nama file."""
    import re
    name = re.sub(r'[^\w\s-]', '', name)
    name = re.sub(r'\s+', '_', name)
    return name.strip("_-")
